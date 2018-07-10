from datetime import datetime
from typing import Callable
import docx
import slidegate
import slg3


FRAME_WIDTHS = [i / 10 for i in range(3, 21)]
GATE_HEIGHTS = [i / 10 for i in range(3, 29, 2)]
HEAD = 6.0
FILE_NAME = 'Drive-power-of-the-slidegates.docx'


def get_surf_slidegate(frame_width: float,
                       gate_height: float) -> slidegate.Slidegate:
    return slidegate.Slidegate(frame_width=frame_width,
                               gate_height=gate_height,
                               hydr_head=gate_height,
                               kind=slidegate.SlgKind.surf,
                               drive=slidegate.Drive.electric,
                               install=slidegate.Install.wall,
                               is_frame_closed=False,
                               have_rack=False,
                               screws_number=1)


def get_deep_slidegate(frame_width: float,
                       gate_height: float) -> slidegate.Slidegate:
    return slidegate.Slidegate(frame_width=frame_width,
                               gate_height=gate_height,
                               hydr_head=HEAD,
                               kind=slidegate.SlgKind.deep,
                               drive=slidegate.Drive.electric,
                               install=slidegate.Install.wall,
                               is_frame_closed=False,
                               have_rack=False,
                               screws_number=1)


def get_flow_slidegate(frame_width: float,
                       gate_height: float) -> slidegate.Slidegate:
    return slidegate.Slidegate(frame_width=frame_width,
                               gate_height=gate_height,
                               hydr_head=gate_height,
                               kind=slidegate.SlgKind.flow,
                               drive=slidegate.Drive.electric,
                               install=slidegate.Install.wall,
                               is_frame_closed=False,
                               have_rack=False,
                               screws_number=1)


def add_power_table(
        doc: docx.Document,
        get_slidegate: Callable[[float, float],
                                slidegate.Slidegate]) -> None:
    table = doc.add_table(rows=len(GATE_HEIGHTS) + 1,
                          cols=len(FRAME_WIDTHS) + 1,
                          style='Medium List 2 Accent 1')
    headers = table.rows[0].cells
    for i, frame_width in enumerate(FRAME_WIDTHS):
        headers[i + 1].text = str(frame_width)
    for i, gate_height in enumerate(GATE_HEIGHTS):
        row = table.rows[i + 1].cells
        row[0].text = str(gate_height)
        for j, frame_width in enumerate(FRAME_WIDTHS):
            slg = get_slidegate(frame_width, gate_height)
            slg3.mass_calculation(slg)
            power = slg.auma.powers[slg.mode]
            weight = slg.mass
            row[j + 1].text = f"{power / 1e3} {weight:.0f}"


def generic_table_powers() -> None:
    doc = docx.Document()

    font = doc.styles['Normal'].font
    font.name = 'Consolas'
    font.size = docx.shared.Pt(10)

    doc.add_paragraph(
        f'Мощности приводов типовых клиновых затворов, кВт ({datetime.now()})'
        '\nОдин невыдвижной винт, настенная установка, без выносной колонки.'
        '\nПо горизонтали ширина рамы, по вертикали высота щита (м).'
        '\nСверху мощность привода (кВт), снизу масса затвора (кг).')

    doc.add_paragraph(f'Поверхностные (напор равен высоте щита):')
    add_power_table(doc, get_surf_slidegate)

    par = doc.add_paragraph()
    par.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    par.add_run(f'Глубинные (напор {HEAD} м. вод. ст.):')
    add_power_table(doc, get_deep_slidegate)

    par = doc.add_paragraph()
    par.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    par.add_run(f'Регулирующие глубинные (напор равен высоте щита):')
    add_power_table(doc, get_flow_slidegate)

    # doc.sections[-1].orientation = docx.enum.section.WD_ORIENTATION.LANDSCAPE
    doc.save(FILE_NAME)


if __name__ == '__main__':
    generic_table_powers()
